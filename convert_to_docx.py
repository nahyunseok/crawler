# -*- coding: utf-8 -*-
"""
KMONG_DESCRIPTION.md → .docx(워드) 변환 스크립트
한글 프로그램에서 열어 .hwp로 저장할 수 있도록 Word 형식으로 변환합니다.
"""
import os
import re
import shutil
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading_elm = cell._tc.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)

def create_kmong_docx():
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ─── 제목 ───
    title = doc.add_heading('', level=0)
    run = title.add_run('🖼️ Gemini 이미지 수집기 V2.0')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('크몽 판매 등록용 서비스 상세 설명서')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('─' * 60)

    # ─── 서비스 제목 ───
    doc.add_heading('✅ 서비스 제목 (예시)', level=2)
    p = doc.add_paragraph()
    run = p.add_run('[이미지 크롤러] 원하는 사이트의 이미지를 한 번에 수집하는 자동화 프로그램 | 엑셀 리포트 | 봇 우회')
    run.font.size = Pt(11)
    run.bold = True

    doc.add_paragraph('')

    # ─── 서비스 설명 본문 ───
    doc.add_heading('✅ 서비스 설명 본문', level=2)

    # 메인 카피
    p = doc.add_paragraph()
    run = p.add_run('🖼️ "원하는 사이트 URL만 넣으면, 이미지를 자동으로 쓸어담습니다"')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    doc.add_paragraph('')
    doc.add_paragraph('이미지를 하나하나 우클릭 → 저장?')
    p = doc.add_paragraph()
    run = p.add_run('이제 그 노가다는 끝입니다.')
    run.bold = True

    doc.add_paragraph('')
    p = doc.add_paragraph('Gemini 이미지 수집기 V2.0은 ')
    run = p.add_run('URL 하나만 입력하면')
    run.bold = True
    p.add_run(' 해당 페이지(+하위 링크)의 모든 이미지를 ')
    run = p.add_run('자동으로 수집 → 다운로드 → 엑셀 리포트까지')
    run.bold = True
    p.add_run(' 한 번에 만들어주는 Windows 전용 프로그램입니다.')

    doc.add_paragraph('─' * 60)

    # ─── 타겟 고객 ───
    doc.add_heading('🔥 이런 분께 딱 맞습니다', level=2)
    targets = [
        '쇼핑몰 상품 이미지를 대량으로 모아야 하는 MD / 마케터',
        '경쟁사 시각 자료를 분석하고 싶은 기획자',
        '블로그·포트폴리오 레퍼런스 이미지를 한꺼번에 저장하고 싶은 디자이너',
        '특정 사이트의 이미지 목록과 메타정보(설명, 출처, 해상도)가 한 파일로 필요한 리서처',
        '"이걸 손으로 하고 있었다고?" 하는 모든 분',
    ]
    for t in targets:
        doc.add_paragraph(f'✔️ {t}')

    doc.add_paragraph('─' * 60)

    # ─── 핵심 기능표 ───
    doc.add_heading('⚡ 핵심 기능 한눈에 보기', level=2)
    features = [
        ('🚀 원클릭 수집', 'URL 입력 → 버튼 한 번이면 끝'),
        ('🛡️ 봇 차단 우회', 'undetected-chromedriver + 랜덤 User-Agent로 사이트 봇 감지를 스마트하게 우회'),
        ('📊 엑셀 리포트', '수집된 이미지마다 파일명·설명·출처·해상도·주변 텍스트를 깔끔한 엑셀(xlsx)로 자동 생성'),
        ('🔍 스마트 필터링', '제외 키워드(logo, banner 등) & 필수 포함 키워드로 원하는 이미지만 골라 수집'),
        ('🔒 로그인 후 수집', '인스타그램·비공개 카페 등 로그인이 필요한 사이트도 지원 (수동 로그인 대기 모드)'),
        ('🔗 딥 크롤링', '현재 페이지뿐 아니라 링크를 타고 한 단계 더 들어가서 수집 (2단계 깊이)'),
        ('📄 자동 페이지 넘기기', '게시판의 "다음" 버튼을 자동 클릭하며 여러 페이지 순회 수집'),
        ('🖼️ 배경 이미지도 수집', '<img> 태그뿐 아니라 CSS background-image, <picture> 태그까지 모두 수집'),
        ('🔄 이어받기(Resume)', '중간에 멈춰도 이전 다운로드 기록을 저장, 재실행 시 중복 없이 이어서 수집'),
        ('⚡ 멀티스레드 다운로드', '5개 동시 다운로드 + 자동 재시도(최대 3회)로 빠르고 안정적'),
        ('👁️ Headless 모드', '브라우저 창을 숨기고 백그라운드에서 조용히 수집 가능'),
        ('📁 자동 정리', '[페이지제목]_도메인_날짜시간 형식으로 폴더 자동 생성'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    hdr = table.rows[0].cells
    hdr[0].text = '기능'
    hdr[1].text = '설명'
    for cell in hdr:
        set_cell_shading(cell, '1A73E8')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    # 데이터
    for feat, desc in features:
        row = table.add_row().cells
        row[0].text = feat
        row[1].text = desc
        # 기능명 굵게
        for p in row[0].paragraphs:
            for run in p.runs:
                run.bold = True

    doc.add_paragraph('─' * 60)

    # ─── 경쟁 비교표 ───
    doc.add_heading('🎯 다른 크롤러와 뭐가 다른가요?', level=2)
    comparisons = [
        ('봇 탐지 우회', '❌ 대부분 차단됨', '✅ undetected-chromedriver + 랜덤 UA'),
        ('로그인 사이트', '❌ 지원 안 됨', '✅ 수동 로그인 대기 모드'),
        ('엑셀 리포트', '❌ 이미지만 저장', '✅ 메타정보(설명, 출처, 해상도, 주변 텍스트) 포함'),
        ('키워드 필터링', '❌ 전부 다 긁어옴', '✅ 제외/필수 키워드, 확장자 필터, 최소 크기 필터'),
        ('배경 이미지', '❌ img 태그만', '✅ CSS background + picture + srcset 모두 수집'),
        ('딥 크롤링', '❌ 단일 페이지', '✅ 2단계 깊이까지 링크 탐색'),
        ('이어받기', '❌ 처음부터 다시', '✅ 중복 방지 히스토리 기반 이어받기'),
        ('안전 딜레이', '❌ 고정 속도', '✅ 랜덤 딜레이 (사람처럼 행동)'),
    ]
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '비교 항목'
    hdr2[1].text = '일반 크롤러 / 확장 프로그램'
    hdr2[2].text = 'Gemini 이미지 수집기 V2.0'
    for cell in hdr2:
        set_cell_shading(cell, '1A73E8')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    for item, other, ours in comparisons:
        row = table2.add_row().cells
        row[0].text = item
        row[1].text = other
        row[2].text = ours
        # 우리 것 셀 배경 연한 초록
        set_cell_shading(row[2], 'E8F5E9')

    doc.add_paragraph('─' * 60)

    # ─── 실행 환경 ───
    doc.add_heading('🖥️ 프로그램 실행 환경', level=2)
    envs = [
        ('운영체제', 'Windows 10 / 11 (64bit)'),
        ('필수 사항', 'Google Chrome 브라우저 설치 (ChromeDriver 자동 관리)'),
        ('형태', 'exe 실행 파일 (설치 불필요, 더블클릭만 하면 됨)'),
        ('라이센스', '1 PC 1 라이선스 (기기 바인딩)'),
    ]
    for label, val in envs:
        p = doc.add_paragraph()
        run = p.add_run(f'• {label}: ')
        run.bold = True
        p.add_run(val)

    doc.add_paragraph('─' * 60)

    # ─── 패키지 ───
    doc.add_heading('📦 패키지별 제공 내역', level=2)

    # Basic
    p = doc.add_paragraph()
    run = p.add_run('▶ BASIC (30,000 ~ 50,000원)')
    run.bold = True
    run.font.size = Pt(12)
    basics = [
        'Gemini 이미지 수집기 V2.0 실행 파일 (.exe)',
        '바탕화면 바로가기 생성 스크립트',
        '상세 사용 매뉴얼 (MANUAL.md)',
        '라이선스 키 1개 (30일)',
        '카카오톡 1:1 기술 지원',
    ]
    for b in basics:
        doc.add_paragraph(f'   ✅ {b}')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('▶ STANDARD (80,000 ~ 120,000원)')
    run.bold = True
    run.font.size = Pt(12)
    standards = [
        'Basic 전체 포함',
        '라이선스 키 1개 (90일)',
        '원격 설치 지원 (팀뷰어/애니데스크)',
        '맞춤 설정 세팅 (수집 대상 사이트에 최적화)',
    ]
    for s in standards:
        doc.add_paragraph(f'   ✅ {s}')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('▶ PREMIUM (150,000 ~ 250,000원)')
    run.bold = True
    run.font.size = Pt(12)
    premiums = [
        'Standard 전체 포함',
        '라이선스 키 1개 (365일)',
        '1:1 사용법 화면 공유 교육 (30분)',
        '업데이트 무상 제공 기간',
    ]
    for pr in premiums:
        doc.add_paragraph(f'   ✅ {pr}')

    doc.add_paragraph('─' * 60)

    # ─── 사용법 ───
    doc.add_heading('⚙️ 이렇게 사용합니다 (3단계)', level=2)
    steps = [
        ('STEP 1.', '프로그램 실행 → 라이선스 키 입력'),
        ('STEP 2.', '수집하고 싶은 사이트 URL 붙여넣기'),
        ('STEP 3.', '[🚀 수집 시작] 클릭 → 끝! (결과 폴더가 자동으로 열립니다)'),
    ]
    for step, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(f'{step} ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        p.add_run(desc)

    doc.add_paragraph('')
    doc.add_paragraph('➡️ 수집 결과 폴더 구조:')
    result_structure = """results/
  └── [페이지제목]_도메인_20260308/
       ├── images/          ← 다운로드된 이미지 파일
       │    ├── 001_photo.jpg
       │    ├── 001_photo.txt  ← 이미지별 메타정보
       │    └── ...
       └── checklist.xlsx    ← 전체 이미지 리포트 (엑셀)"""
    p = doc.add_paragraph()
    run = p.add_run(result_structure)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('─' * 60)

    # ─── 보안 시스템 ───
    doc.add_heading('🔐 보안 & 라이선스 시스템', level=2)
    securities = [
        ('기기 바인딩', 'MAC + 하드웨어 정보 기반 고유 ID로 1PC 1라이선스 보장'),
        ('온라인 인증', '구글 스프레드시트 기반 실시간 키 검증'),
        ('오프라인 대응', '인터넷 없어도 캐시 기반으로 기존 키 유지'),
        ('Brute Force 방지', '5회 연속 실패 시 1분간 입력 차단'),
        ('이용약관 동의', '첫 실행 시 면책 조항 동의 필수'),
    ]
    for label, val in securities:
        p = doc.add_paragraph()
        run = p.add_run(f'• {label}: ')
        run.bold = True
        p.add_run(val)

    doc.add_paragraph('─' * 60)

    # ─── FAQ ───
    doc.add_heading('❓ 자주 묻는 질문 (FAQ)', level=2)
    faqs = [
        ('Q. Mac에서도 되나요?', 'A. 현재 Windows 전용입니다. Mac 버전은 추후 업데이트 예정입니다.'),
        ('Q. 크롬 드라이버를 따로 설치해야 하나요?', 'A. 아닙니다. 프로그램이 자동으로 호환되는 드라이버를 관리합니다.'),
        ('Q. 어떤 사이트든 수집 가능한가요?', 'A. 대부분의 일반 웹사이트에서 작동합니다. 단, 각 사이트의 이용약관과 robots.txt 정책을 반드시 확인하고 준수해야 합니다.'),
        ('Q. 수집 중에 에러가 나면?', 'A. "안전 딜레이"를 높이면 대부분 해결됩니다. 수집 도중 멈추더라도 이미 다운로드된 파일은 보존되며, 이어받기도 가능합니다.'),
        ('Q. 라이선스 기간이 끝나면?', 'A. 프로그램 내 [라이선스 갱신] 버튼으로 언제든 연장할 수 있습니다.'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        doc.add_paragraph(a)
        doc.add_paragraph('')

    doc.add_paragraph('─' * 60)

    # ─── 면책 조항 ───
    doc.add_heading('⚠️ 면책 조항 (필수 고지)', level=2)
    p = doc.add_paragraph()
    run = p.add_run('본 프로그램의 사용으로 인해 발생하는 모든 문제(계정 정지, IP 차단, 지적 재산권 침해, 금전적 손실 등)에 대한 책임은 전적으로 사용자 본인에게 있습니다. ')
    run.bold = True
    p.add_run('타인의 저작물을 무단으로 수집·배포하는 행위는 관련 법률에 의해 처벌될 수 있으니, 대상 사이트의 이용약관과 robots.txt를 반드시 확인하십시오.')

    doc.add_paragraph('─' * 60)

    # ─── 등록 팁 ───
    doc.add_heading('🏷️ 크몽 등록 시 추천 태그', level=2)
    doc.add_paragraph('이미지 크롤러, 웹 크롤링, 이미지 수집기, 자동화 프로그램, 크롤링 프로그램, 이미지 다운로더, 웹 스크래핑, 엑셀 리포트, 대량 이미지 수집, 봇 우회')

    doc.add_paragraph('')
    doc.add_heading('📋 크몽 썸네일 문구 제안', level=2)
    p = doc.add_paragraph()
    run = p.add_run('메인 카피: ')
    run.bold = True
    doc.add_paragraph('🖼️ 이미지, 아직도 하나씩 저장하세요? URL 하나로 수백 장 자동 수집!')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('서브 카피: ')
    run.bold = True
    doc.add_paragraph('✅ 봇 차단 우회  ✅ 엑셀 리포트  ✅ 로그인 사이트 지원  ✅ 딥 크롤링  ✅ 키워드 필터  ✅ exe 바로 실행')

    # ─── 저장 ───
    desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
    output_path = os.path.join(desktop, 'Gemini_이미지_수집기_크몽_설명서.docx')
    doc.save(output_path)
    print(f'✅ 파일 저장 완료: {output_path}')
    print(f'💡 한글(.hwp)로 변환하려면: 한글 프로그램에서 이 파일을 열고 → 다른 이름으로 저장 → .hwp 선택')

if __name__ == '__main__':
    create_kmong_docx()
